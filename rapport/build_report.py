"""
Build the ISJ (Licence 2) academic internship report as a Microsoft Word (.docx) file.

Report language: English (as requested by the student).
Structure follows the "Guide de Rédaction du Rapport de Stage ISJ-2025".

Run:
    python3 rapport/build_report.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from content import write_content

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(PROJECT_ROOT, "frontend", "public", "Events")
UML_PNG = {
    "class": os.path.join(FIG_DIR, "ClassDaigram.png"),
    "usecase": os.path.join(FIG_DIR, "useCase.png"),
    "ticket": os.path.join(FIG_DIR, "sequenceForTicket.png"),
    "qr": os.path.join(FIG_DIR, "sequenceForQRScann.png"),
}

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "Rapport_de_Stage_DODGE_Culhane.docx")

BLACK = RGBColor(0x00, 0x00, 0x00)
BURGUNDY = RGBColor(0x57, 0x00, 0x13)


# --------------------------------------------------------------------------- #
# Low level helpers
# --------------------------------------------------------------------------- #
def set_base_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading styles: bold, black, Times New Roman, no colour, modest sizes.
    for name, size in (
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLACK
        st.font.italic = False
        sp = st.paragraph_format
        sp.space_before = Pt(12)
        sp.space_after = Pt(6)
        sp.keep_with_next = True


def _field(paragraph, instr):
    """Insert a Word field (e.g. PAGE, TOC) into a paragraph."""
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose \"Update Field\" to build the list."
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(t)
    run._r.append(fld2)
    # mark dirty so Word/LibreOffice prompts for update on open
    dirty = OxmlElement("w:dirty")
    dirty.set(qn("w:val"), "true")
    instr_text.addprevious(dirty)


def set_margins(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_page_numbering(section, fmt, restart_at=None):
    """Set the page-number format ('decimal' / 'upperRoman') and optional
    restart value directly on the section's sectPr (python-docx 1.2 does not
    expose page_number_format)."""
    sectPr = section._sectPr
    pg_num = sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sectPr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if restart_at is not None:
        pg_num.set(qn("w:start"), str(restart_at))


def add_page_number_footer(section, fmt, restart_at=None):
    set_page_numbering(section, fmt, restart_at)
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    _field(paragraph, "PAGE")


# --------------------------------------------------------------------------- #
# High level "H" builder API consumed by content.py
# --------------------------------------------------------------------------- #
class Report:
    def __init__(self, doc):
        self.doc = doc
        self.figure_no = 0
        self.table_no = 0

    def title_block(self, lines, subtitle=None):
        for ln in lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(ln)
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = BURGUNDY
        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(subtitle)
            r.bold = True
            r.font.size = Pt(13)

    def h1(self, text):
        self.doc.add_heading(text, level=1)

    def h2(self, text):
        self.doc.add_heading(text, level=2)

    def h3(self, text):
        self.doc.add_heading(text, level=3)

    def p(self, text):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Cm(1)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return para

    def bullet(self, text):
        para = self.doc.add_paragraph(style="List Bullet")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return para

    def centered(self, text, bold=False, size=12, color=None):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        return para

    def spacer(self):
        self.doc.add_paragraph()

    def pagebreak(self):
        self.doc.add_page_break()

    def sectionbreak(self):
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    def figure(self, path, caption):
        self.figure_no += 1
        if os.path.exists(path):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(path, width=Cm(15))
        else:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run("[ Figure file not found: %s ]" % os.path.basename(path))
            r.italic = True
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Figure %d: %s" % (self.figure_no, caption))
        r.italic = True
        r.font.size = Pt(11)
        self.spacer()

    def table(self, headers, rows, caption, source=None):
        self.table_no += 1
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(str(val))
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Table %d: %s" % (self.table_no, caption))
        r.italic = True
        r.font.size = Pt(11)
        if source:
            src = self.doc.add_paragraph()
            src.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rs = src.add_run("Source: %s" % source)
            rs.italic = True
            rs.font.size = Pt(10)
        self.spacer()

    def toc(self, levels="1-2", title="Table of Contents"):
        self.h1(title) if title else None
        para = self.doc.add_paragraph()
        _field(para, 'TOC \\o "%s" \\h \\z \\u' % levels)


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def main():
    doc = Document()
    set_base_style(doc)
    set_margins(doc.sections[0])

    # Front matter -> roman numerals
    add_page_number_footer(doc.sections[0], "upperRoman")

    H = Report(doc)
    write_content(H, UML_PNG)

    # write_content() performed one section break before the General
    # Introduction, creating a second section. Configure it: arabic numerals,
    # restart at 1, same margins.
    body = doc.sections[-1]
    set_margins(body)
    add_page_number_footer(
        body, "decimal", restart_at=1
    )

    doc.save(OUT_PATH)
    print("Saved:", OUT_PATH)


if __name__ == "__main__":
    main()
